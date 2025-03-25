from pathlib import Path
from typing import Dict, Any, List
import logging
from docx import Document
from processors.base import BaseProcessor

logger = logging.getLogger(__name__)

class DOCXProcessor(BaseProcessor):
    """Processor for DOCX documents."""
    
    def __init__(self):
        super().__init__()
        self.supported_extensions = ['.docx']
    
    async def extract_metadata(self, file_path: Path) -> Dict[str, Any]:
        """Extract metadata from DOCX document."""
        try:
            metadata = self._generate_metadata(file_path)
            
            doc = Document(file_path)
            core_properties = doc.core_properties
            
            metadata.update({
                "title": core_properties.title or '',
                "author": core_properties.author or '',
                "subject": core_properties.subject or '',
                "keywords": core_properties.keywords or '',
                "created": core_properties.created.isoformat() if core_properties.created else '',
                "modified": core_properties.modified.isoformat() if core_properties.modified else '',
                "last_modified_by": core_properties.last_modified_by or '',
                "revision": core_properties.revision or 0,
                "section_count": len(doc.sections)
            })
            
            return metadata
            
        except Exception as e:
            logger.error(f"Error extracting DOCX metadata: {str(e)}")
            return self._generate_metadata(file_path)
    
    async def extract_text(self, file_path: Path) -> str:
        """Extract text content from DOCX document."""
        try:
            doc = Document(file_path)
            text_content = []
            
            # Extract text from paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    text_content.append(para.text)
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text)
                    if row_text:
                        text_content.append(" | ".join(row_text))
            
            return "\n".join(text_content)
            
        except Exception as e:
            logger.error(f"Error extracting DOCX text: {str(e)}")
            return ""
    
    async def extract_tables(self, file_path: Path) -> List[Dict[str, Any]]:
        """Extract tables from DOCX document."""
        try:
            tables = []
            doc = Document(file_path)
            
            for table_index, table in enumerate(doc.tables, 1):
                table_data = {
                    "table_id": f"table_{table_index}",
                    "rows": [],
                    "column_count": len(table.columns),
                    "row_count": len(table.rows)
                }
                
                for row in table.rows:
                    row_data = []
                    for cell in row.cells:
                        row_data.append(cell.text.strip())
                    table_data["rows"].append(row_data)
                
                tables.append(table_data)
            
            return tables
            
        except Exception as e:
            logger.error(f"Error extracting DOCX tables: {str(e)}")
            return []
    
    async def extract_charts(self, file_path: Path) -> List[Dict[str, Any]]:
        """Extract charts/graphs from DOCX document."""
        # Note: This is a placeholder. For actual chart extraction,
        # you would need to use additional libraries to handle embedded charts
        try:
            charts = []
            # TODO: Implement chart extraction from DOCX
            return charts
        except Exception as e:
            logger.error(f"Error extracting DOCX charts: {str(e)}")
            return [] 