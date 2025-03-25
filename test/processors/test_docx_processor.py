import asyncio
import os
from pathlib import Path
import pytest
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from processors.factory import ProcessorFactory

# Create test data directory if it doesn't exist
TEST_DATA_DIR = Path(__file__).parent / "test_data"
TEST_DATA_DIR.mkdir(exist_ok=True)

def create_test_docx():
    """Create a test DOCX file with sample content."""
    doc = Document()
    
    # Add title
    title = doc.add_heading('Investment Analysis Report', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add some metadata through core properties
    doc.core_properties.author = "Financial Analyst"
    doc.core_properties.title = "Investment Analysis Report"
    doc.core_properties.subject = "Quarterly Financial Review"
    doc.core_properties.keywords = "investment, finance, quarterly report"
    
    # Add a paragraph with company information
    doc.add_paragraph('Company: XYZ Investments\nDate: January 1, 2024\n')
    
    # Add a heading for the executive summary
    doc.add_heading('Executive Summary', level=1)
    summary = doc.add_paragraph()
    summary.add_run('This report provides a comprehensive analysis of our investment portfolio performance for Q4 2023. ')
    summary.add_run('Key highlights include strong returns in technology sector and emerging markets.')
    
    # Add a table with investment data
    doc.add_heading('Portfolio Performance', level=1)
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    
    # Add header row
    header_cells = table.rows[0].cells
    headers = ['Asset Class', 'Allocation', 'Return', 'Risk Score']
    for i, header in enumerate(headers):
        header_cells[i].text = header
    
    # Add data rows
    data = [
        ['Equities', '60%', '12.5%', 'High'],
        ['Fixed Income', '30%', '4.2%', 'Low'],
        ['Alternatives', '10%', '8.7%', 'Medium']
    ]
    
    for row_data in data:
        row_cells = table.add_row().cells
        for i, value in enumerate(row_data):
            row_cells[i].text = value
    
    # Add a section with bullet points
    doc.add_heading('Key Findings', level=1)
    findings = [
        'Technology sector outperformed with 15% returns',
        'Fixed income provided stable returns despite market volatility',
        'ESG investments showed strong momentum',
        'Emerging markets presented new opportunities'
    ]
    
    for finding in findings:
        doc.add_paragraph(finding, style='List Bullet')
    
    # Add a recommendations section
    doc.add_heading('Recommendations', level=1)
    recommendations = doc.add_paragraph()
    recommendations.add_run('Based on our analysis, we recommend:').bold = True
    doc.add_paragraph('1. Increase technology sector exposure')
    doc.add_paragraph('2. Maintain current fixed income allocation')
    doc.add_paragraph('3. Explore emerging market opportunities')
    
    # Save the document
    test_file = TEST_DATA_DIR / "test_analysis.docx"
    doc.save(test_file)
    return test_file

@pytest.mark.asyncio
async def test_docx_processor():
    """Test the DOCX processor with a sample file."""
    # Create test DOCX file
    test_file = create_test_docx()
    
    try:
        # Get the DOCX processor
        processor = ProcessorFactory.get_processor(test_file)
        
        # Process the document
        result = await processor.process(test_file)
        
        # Verify processing status
        assert result["processing_status"] == "success"
        
        # Verify metadata
        metadata = result["metadata"]
        assert metadata["extension"] == ".docx"
        assert metadata["title"] == "Investment Analysis Report"
        assert metadata["author"] == "Financial Analyst"
        assert metadata["subject"] == "Quarterly Financial Review"
        assert "investment" in metadata["keywords"].lower()
        
        # Verify content
        content = result["content"]
        
        # Check text content
        assert len(content["text"]) > 0
        assert "XYZ Investments" in content["text"]
        assert "Executive Summary" in content["text"]
        assert "Portfolio Performance" in content["text"]
        
        # Verify tables
        assert len(content["tables"]) > 0
        table = content["tables"][0]  # First table
        
        # Check table structure
        assert table["table_id"] == "table_1"
        assert len(table["rows"]) == 4  # Header + 3 data rows
        
        # Check first row (headers)
        first_row = table["rows"][0]
        assert "Asset Class" in first_row
        assert "Allocation" in first_row
        assert "Return" in first_row
        assert "Risk Score" in first_row
        
        # Check second row (first data row)
        second_row = table["rows"][1]
        assert "Equities" in second_row
        assert "60%" in second_row
        assert "12.5%" in second_row
        assert "High" in second_row
        
        print("\nTest Results:")
        print("-" * 50)
        print(f"File: {test_file}")
        print(f"Status: {result['processing_status']}")
        print("\nMetadata:")
        for key, value in metadata.items():
            print(f"  {key}: {value}")
        
        print("\nContent Summary:")
        print(f"  Text Length: {len(content['text'])} characters")
        print(f"  Tables Found: {len(content['tables'])}")
        
        print("\nFirst Table Structure:")
        print(f"  Total Rows: {len(table['rows'])}")
        print(f"  First Row: {table['rows'][0]}")
        print(f"  Second Row: {table['rows'][1]}")
        
        # Print first 200 characters of extracted text
        print("\nExtracted Text Preview:")
        print(f"  {content['text'][:200]}...")
        
    finally:
        # Clean up test file
        if test_file.exists():
            test_file.unlink()

if __name__ == "__main__":
    asyncio.run(test_docx_processor()) 